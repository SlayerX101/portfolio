from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Oageng_James_Tsumaki_Junior_CSharp_Developer_CV.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            set_cell_margins(row.cells[index])
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text.upper())
    run.bold = True
    return p


def add_role(doc, title, organization, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(24, 50, 74)
    p.add_run(f" | {organization}")
    for bullet in bullets:
        item = doc.add_paragraph(style="List Bullet")
        item.add_run(bullet)


def add_project(doc, title, summary, bullets, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(24, 50, 74)
    doc.add_paragraph(summary)
    for bullet in bullets:
        item = doc.add_paragraph(style="List Bullet")
        item.add_run(bullet)
    tech_p = doc.add_paragraph()
    tech_p.paragraph_format.space_after = Pt(4)
    tech_run = tech_p.add_run(f"Tech: {tech}")
    tech_run.bold = True
    tech_run.font.color.rgb = RGBColor(47, 125, 117)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(23, 33, 47)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    heading = styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(12.5)
    heading.font.color.rgb = RGBColor(46, 116, 181)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)

    for style_name in ["List Bullet"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.12)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Oageng James Tsumaki")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(24, 50, 74)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    sub_run = subtitle.add_run("Junior C# Developer | ASP.NET Core Web API | Software Support")
    sub_run.bold = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(47, 125, 117)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(8)
    contact.add_run("Pretoria East, South Africa | 081 042 0348 | oagengtsumakijr@gmail.com")

    add_heading(doc, "Professional Profile")
    doc.add_paragraph(
        "Bachelor of Software Engineering graduate focused on junior C# developer, backend developer, and software support opportunities. "
        "Practical experience includes ASP.NET Core Web API portfolio projects, responsive web development, JavaScript fundamentals, digital booking systems, "
        "Excel-based data handling, and customer-facing troubleshooting. Comfortable building REST APIs, validating requests, applying business rules, working with structured data, "
        "and communicating clearly with users and teams."
    )

    add_heading(doc, "Key Skills")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.55, 5.45])
    skills = [
        ("Programming", "C#, ASP.NET Core Web API, .NET Minimal APIs, HTML, CSS, JavaScript, responsive web design, SQL fundamentals"),
        ("Backend", "REST endpoints, CRUD operations, route design, request validation, JSON, LINQ, records, collections"),
        ("Data and Tools", "Microsoft Excel, Word, PowerPoint, Outlook, Git, GitHub, Visual Studio Code, Postman or REST Client"),
        ("Professional", "Problem solving, communication, adaptability, teamwork, client feedback management, troubleshooting mindset"),
    ]
    for row_index, (label, detail) in enumerate(skills):
        cells = table.rows[row_index].cells
        set_cell_shading(cells[0], "E8EEF5")
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(detail)

    add_heading(doc, "Portfolio Projects")
    add_project(
        doc,
        "Service Desk Pro API",
        "Advanced ticketing API with SLA calculations, assignment workflow, comments, status transitions, breached-ticket filters, audit trail, and dashboard metrics.",
        ["Implemented service-layer business rules for ticket workflow and SLA due times.", "Added audit events and dashboard reporting for operational visibility."],
        "C#, ASP.NET Core, Minimal APIs, DTOs, service layer, business rules",
    )
    add_project(
        doc,
        "Inventory Orders API",
        "Inventory and order management API with product search, stock adjustments, order placement, VAT totals, cancellation stock release, low-stock reporting, and sales analytics.",
        ["Built order validation, stock reservation, and cancellation stock release rules.", "Created low-stock and sales report endpoints using LINQ summaries."],
        "C#, ASP.NET Core, REST, validation, reporting",
    )
    add_project(
        doc,
        "Learning Progress API",
        "Learning platform API with courses, enrollments, module completion, quiz scoring, pass rules, learner dashboards, certificate readiness, and course performance reports.",
        ["Modeled course modules, enrollments, quiz submissions, and learner progress.", "Calculated completion percentages and course performance analytics."],
        "C#, ASP.NET Core, LINQ, analytics, domain modeling",
    )
    add_project(
        doc,
        "Task Tracker API",
        "C# Web API for managing tasks with statuses, priorities, due dates, filtering, updates, deletion, and summary reporting.",
        ["Built CRUD endpoints for task management.", "Used LINQ for filtering and grouped status summaries."],
        "C#, ASP.NET Core, Minimal APIs, LINQ",
    )
    add_project(
        doc,
        "Expense Tracker API",
        "Budgeting API that records expenses and returns monthly totals, averages, and category-based reports.",
        ["Validated required fields and positive money values.", "Used decimal calculations for practical finance-style data."],
        "C#, ASP.NET Core, REST, LINQ",
    )
    add_project(
        doc,
        "Library API",
        "Small library API for books, members, active loans, returns, and availability rules.",
        ["Modeled books, members, and loans with clear records.", "Added business rules to prevent unavailable book loans."],
        "C#, ASP.NET Core, Minimal APIs",
    )
    add_project(
        doc,
        "Job Application Tracker API",
        "Job-search workflow API for roles, companies, statuses, follow-up dates, notes, and application statistics.",
        ["Created status update and follow-up routes.", "Grouped application counts for dashboard-style summaries."],
        "C#, ASP.NET Core, LINQ, JSON",
    )
    add_project(
        doc,
        "Weather Journal API",
        "Weather logging API for city observations, humidity checks, temperature averages, and warmest-entry summaries.",
        ["Added city/date filters and basic input validation.", "Calculated simple analytics from stored entries."],
        "C#, ASP.NET Core, Minimal APIs",
    )

    add_heading(doc, "Relevant Experience")
    add_role(
        doc,
        "Contact Centre Agent",
        "Bounce Inc | Current",
        [
            "Use digital booking systems to process customer requests and troubleshoot booking issues.",
            "Capture and manage client information accurately in a fast-paced service environment.",
            "Communicate technical and process-related information clearly to customers and team members.",
            "Complete sales and appointment-setting targets while maintaining professional service quality.",
        ],
    )
    add_role(
        doc,
        "Customer Service & Safety Coach",
        "Bounce Inc | Current",
        [
            "Support front desk operations, customer guidance, and system-based booking management.",
            "Help customers solve issues calmly while following safety and operating procedures.",
            "Develop communication habits useful for IT support and user-facing technical roles.",
        ],
    )
    add_role(
        doc,
        "Data Capturer",
        "Tsumaki Electricians | Since 2020",
        [
            "Manage structured invoice and company data using Excel and digital filing processes.",
            "Maintain accurate records and support basic data organization for business administration.",
            "Build practical understanding of small-business operations and documentation workflows.",
        ],
    )
    add_role(
        doc,
        "Freelance Web Developer",
        "Independent Projects | Project-Based",
        [
            "Design and develop responsive websites using HTML, CSS, and JavaScript.",
            "Translate client requirements into clean layouts, mobile-friendly pages, and simple user journeys.",
            "Manage revisions and client feedback professionally.",
        ],
    )

    add_heading(doc, "Education")
    education = doc.add_paragraph()
    education.add_run("Bachelor of Software Engineering").bold = True
    education.add_run(" | Eduvos | Completed, 2023-2025")
    school = doc.add_paragraph()
    school.add_run("National Senior Certificate - Bachelor's Pass").bold = True
    school.add_run(" | Willowridge High School | 2021")
    doc.add_paragraph("Subjects: Mathematics, Physical Sciences, Information Technology")

    add_heading(doc, "Certifications")
    for cert in [
        "Web Design 101 - Webflow",
        "Digital Marketing Certificate - Google, 2023",
        "Emergency First Aid Level 1, 2025",
    ]:
        doc.add_paragraph(cert, style="List Bullet")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
